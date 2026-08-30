#!/usr/bin/env python3
"""Safely sync a controlled set of public WPA Zenodo records into Cloudflare R2.

Design goals:
- additive/idempotent: never deletes R2 objects;
- provenance-preserving: every chunk carries Zenodo record/file metadata;
- completion marker per record: metadata JSON is uploaded only after all chunks;
- strict verification: every target record must yield at least one text chunk;
- bounded chunks suitable for AI Search ingestion.
"""
from __future__ import annotations

import argparse
import hashlib
import html
import io
import json
import os
import re
import sys
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests
from docx import Document
from pypdf import PdfReader

ZENODO_API = "https://zenodo.org/api/records"
DEFAULT_BUCKET = "protocol-kb"
DEFAULT_PREFIX = (
    "world-protocol-academy/UPLOAD_TO_WORLD_PROTOCOL_ACADEMY/"
    "99_ai_search_ready"
)
CHUNK_CHARS = 24000
OVERLAP_CHARS = 1200
MIN_EXTRACTED_CHARS = 500
MAX_SOURCE_BYTES = 80 * 1024 * 1024
USER_AGENT = "WorldProtocolAcademy-Zenodo-R2-Sync/1.0 (+https://worldprotocolacademy.mk/)"


@dataclass
class SourceFile:
    name: str
    url: str
    size: int | None = None
    checksum: str | None = None


def log(msg: str) -> None:
    print(msg, flush=True)


def clean_html(value: str) -> str:
    value = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", value or "")
    value = re.sub(r"(?i)<br\s*/?>", "\n", value)
    value = re.sub(r"(?i)</p\s*>", "\n\n", value)
    value = re.sub(r"<[^>]+>", " ", value)
    value = html.unescape(value)
    value = value.replace("\xa0", " ")
    return normalize_text(value)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def slugify(value: str, max_len: int = 90) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return (value[:max_len].rstrip("-") or "source")


def request_json(session: requests.Session, url: str, **kwargs: Any) -> dict[str, Any]:
    for attempt in range(1, 5):
        try:
            r = session.get(url, timeout=45, **kwargs)
            r.raise_for_status()
            return r.json()
        except Exception as exc:
            if attempt == 4:
                raise
            log(f"WARN retry {attempt}/3 for {url}: {exc}")
            time.sleep(attempt * 2)
    raise RuntimeError("unreachable")


def zenodo_record(session: requests.Session, record_id: str) -> dict[str, Any]:
    return request_json(session, f"{ZENODO_API}/{record_id}")


def creators_from_metadata(meta: dict[str, Any]) -> list[str]:
    out: list[str] = []
    for c in meta.get("creators") or []:
        if isinstance(c, str):
            out.append(c)
            continue
        if not isinstance(c, dict):
            continue
        po = c.get("person_or_org")
        if isinstance(po, dict) and po.get("name"):
            out.append(str(po["name"]))
        elif c.get("name"):
            out.append(str(c["name"]))
    return out


def extract_doi(record: dict[str, Any]) -> str:
    pids = record.get("pids") or {}
    if isinstance(pids, dict):
        doi = pids.get("doi")
        if isinstance(doi, dict) and doi.get("identifier"):
            return str(doi["identifier"])
        if isinstance(doi, str):
            return doi
    meta = record.get("metadata") or {}
    for key in ("doi", "prereserve_doi"):
        val = meta.get(key)
        if isinstance(val, str):
            return val
        if isinstance(val, dict) and val.get("doi"):
            return str(val["doi"])
    return ""


def resource_type(meta: dict[str, Any]) -> str:
    rt = meta.get("resource_type") or meta.get("upload_type") or ""
    if isinstance(rt, str):
        return rt
    if isinstance(rt, dict):
        return str(rt.get("title") or rt.get("id") or rt.get("type") or "")
    return ""


def get_file_entries(session: requests.Session, record: dict[str, Any]) -> list[SourceFile]:
    files = record.get("files")
    raw_entries: list[dict[str, Any]] = []
    if isinstance(files, list):
        raw_entries = [x for x in files if isinstance(x, dict)]
    elif isinstance(files, dict):
        entries = files.get("entries")
        if isinstance(entries, dict):
            raw_entries = [x for x in entries.values() if isinstance(x, dict)]
        elif isinstance(entries, list):
            raw_entries = [x for x in entries if isinstance(x, dict)]

    if not raw_entries:
        files_url = (record.get("links") or {}).get("files")
        if files_url:
            payload = request_json(session, str(files_url))
            entries = payload.get("entries") if isinstance(payload, dict) else None
            if isinstance(entries, dict):
                raw_entries = [x for x in entries.values() if isinstance(x, dict)]
            elif isinstance(entries, list):
                raw_entries = [x for x in entries if isinstance(x, dict)]

    out: list[SourceFile] = []
    for e in raw_entries:
        name = str(e.get("key") or e.get("filename") or e.get("name") or "").strip()
        links = e.get("links") or {}
        url = ""
        if isinstance(links, dict):
            url = str(links.get("content") or links.get("download") or links.get("self") or "")
        if not name or not url:
            continue
        size = e.get("size")
        try:
            size_i = int(size) if size is not None else None
        except Exception:
            size_i = None
        checksum = e.get("checksum")
        out.append(SourceFile(name=name, url=url, size=size_i, checksum=str(checksum) if checksum else None))
    return out


def file_score(f: SourceFile) -> tuple[int, int, str]:
    n = f.name.lower()
    ext = Path(n).suffix
    ext_score = {".pdf": 500, ".docx": 400, ".md": 300, ".txt": 250}.get(ext, 0)
    positive = sum(
        w for token, w in {
            "final": 90,
            "publication": 60,
            "paper": 50,
            "protocol": 45,
            "working": 40,
            "bilingual": 85,
            "english": 30,
            "macedonian": 30,
            "mk": 10,
            "en": 10,
        }.items() if token in n
    )
    negative = sum(
        w for token, w in {
            "qa": 180,
            "review": 180,
            "report": 30,
            "metadata": 220,
            "readme": 220,
            "citation": 220,
            "bib": 220,
            "figure": 160,
            "cover": 160,
            "supplement": 140,
            "annex": 80,
        }.items() if token in n
    )
    size_penalty = 80 if (f.size or 0) > MAX_SOURCE_BYTES else 0
    return (ext_score + positive - negative - size_penalty, -(f.size or 0), f.name)


def choose_source_files(entries: list[SourceFile]) -> list[SourceFile]:
    eligible = [f for f in entries if Path(f.name.lower()).suffix in {".pdf", ".docx", ".md", ".txt"}]
    if not eligible:
        return []

    blocked_tokens = ("qa", "review", "metadata", "readme", "citation", "bibliography", "figure", "cover", "supplement")
    primary = [f for f in eligible if not any(t in f.name.lower() for t in blocked_tokens)]
    pool = primary or eligible
    pool.sort(key=file_score, reverse=True)

    # Prefer PDFs. Keep at most two likely primary files so separate MK/EN editions
    # can both enter the corpus, without indexing QA/supplementary material.
    pdfs = [f for f in pool if Path(f.name.lower()).suffix == ".pdf" and file_score(f)[0] >= 300]
    if pdfs:
        return pdfs[:2]
    return pool[:2]


def download_file(session: requests.Session, source: SourceFile) -> tuple[bytes, str]:
    with session.get(source.url, timeout=120, stream=True) as r:
        r.raise_for_status()
        buf = io.BytesIO()
        h = hashlib.sha256()
        total = 0
        for part in r.iter_content(chunk_size=1024 * 1024):
            if not part:
                continue
            total += len(part)
            if total > MAX_SOURCE_BYTES:
                raise RuntimeError(f"source file exceeds {MAX_SOURCE_BYTES} bytes: {source.name}")
            h.update(part)
            buf.write(part)
        return buf.getvalue(), h.hexdigest()


def extract_text(name: str, data: bytes) -> str:
    ext = Path(name.lower()).suffix
    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(data), strict=False)
        pages: list[str] = []
        for i, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:
                log(f"WARN PDF page extraction failed {name} p{i}: {exc}")
                text = ""
            if text.strip():
                pages.append(f"\n\n[Page {i}]\n{text}")
        return normalize_text("".join(pages))
    if ext == ".docx":
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return normalize_text("\n\n".join(parts))
    if ext in {".md", ".txt"}:
        return normalize_text(data.decode("utf-8", errors="replace"))
    return ""


def chunk_text(text: str, size: int = CHUNK_CHARS, overlap: int = OVERLAP_CHARS) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        target_end = min(n, start + size)
        end = target_end
        if target_end < n:
            floor = start + max(size // 2, 1000)
            cut = text.rfind("\n\n", floor, target_end)
            if cut < floor:
                cut = text.rfind("\n", floor, target_end)
            if cut >= floor:
                end = cut
        part = text[start:end].strip()
        if part:
            chunks.append(part)
        if end >= n:
            break
        next_start = max(end - overlap, start + 1)
        start = next_start
    return chunks


class R2Client:
    def __init__(self, account_id: str, api_token: str, bucket: str, session: requests.Session):
        self.account_id = account_id
        self.api_token = api_token
        self.bucket = bucket
        self.session = session
        self.base = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/r2/buckets/{bucket}/objects"
        self.headers = {"Authorization": f"Bearer {api_token}"}

    def list_keys(self, prefix: str) -> dict[str, int]:
        keys: dict[str, int] = {}
        cursor: str | None = None
        while True:
            params: dict[str, Any] = {"prefix": prefix, "per_page": 1000}
            if cursor:
                params["cursor"] = cursor
            r = self.session.get(self.base, headers=self.headers, params=params, timeout=60)
            r.raise_for_status()
            payload = r.json()
            if payload.get("success") is not True:
                raise RuntimeError(f"Cloudflare list failed: {payload}")
            for obj in payload.get("result") or []:
                key = str(obj.get("key") or "")
                if key:
                    keys[key] = int(obj.get("size") or 0)
            info = payload.get("result_info") or {}
            if not info.get("is_truncated"):
                break
            cursor = info.get("cursor")
            if not cursor:
                raise RuntimeError("Cloudflare list was truncated without cursor")
        return keys

    def put(self, key: str, body: bytes, content_type: str) -> None:
        url = f"{self.base}/{key}"
        headers = dict(self.headers)
        headers["Content-Type"] = content_type
        r = self.session.put(url, headers=headers, data=body, timeout=120)
        r.raise_for_status()
        try:
            payload = r.json()
        except Exception:
            payload = None
        if isinstance(payload, dict) and payload.get("success") is False:
            raise RuntimeError(f"Cloudflare upload failed for {key}: {payload}")


def metadata_for_record(record_id: str, record: dict[str, Any]) -> dict[str, Any]:
    meta = record.get("metadata") or {}
    title = str(meta.get("title") or "").strip()
    desc = clean_html(str(meta.get("description") or meta.get("notes") or ""))
    links = record.get("links") or {}
    return {
        "schema_version": "1.0",
        "source": "Zenodo",
        "zenodo_record_id": record_id,
        "zenodo_url": str(links.get("self_html") or links.get("html") or f"https://zenodo.org/records/{record_id}"),
        "doi": extract_doi(record),
        "title": title,
        "creators": creators_from_metadata(meta),
        "publication_date": str(meta.get("publication_date") or meta.get("date") or ""),
        "resource_type": resource_type(meta),
        "description": desc,
    }


def yaml_quote(s: str) -> str:
    return json.dumps(s, ensure_ascii=False)


def chunk_document_header(meta: dict[str, Any], source: SourceFile, source_sha256: str, chunk_no: int, chunk_total: int) -> str:
    creators = "; ".join(meta.get("creators") or [])
    lines = [
        "---",
        "wpa_schema: \"zenodo-ai-ready-v1\"",
        "source: \"Zenodo\"",
        f"zenodo_record_id: {yaml_quote(str(meta['zenodo_record_id']))}",
        f"doi: {yaml_quote(str(meta.get('doi') or ''))}",
        f"title: {yaml_quote(str(meta.get('title') or ''))}",
        f"creators: {yaml_quote(creators)}",
        f"publication_date: {yaml_quote(str(meta.get('publication_date') or ''))}",
        f"resource_type: {yaml_quote(str(meta.get('resource_type') or ''))}",
        f"zenodo_url: {yaml_quote(str(meta.get('zenodo_url') or ''))}",
        f"source_file: {yaml_quote(source.name)}",
        f"source_file_sha256: {yaml_quote(source_sha256)}",
        f"source_file_checksum_zenodo: {yaml_quote(source.checksum or '')}",
        f"chunk_number: {chunk_no}",
        f"chunk_total: {chunk_total}",
        "---",
        "",
        f"# {meta.get('title') or 'Untitled'}",
        "",
    ]
    return "\n".join(lines)


def sync_record(
    session: requests.Session,
    r2: R2Client,
    base_prefix: str,
    record_id: str,
    existing: dict[str, int],
) -> dict[str, Any]:
    marker = f"{base_prefix}/_metadata/zenodo/{record_id}.json"
    if marker in existing:
        log(f"SKIP {record_id}: completion marker already exists")
        return {"record_id": record_id, "status": "skipped_existing", "objects": 0}

    record = zenodo_record(session, record_id)
    meta = metadata_for_record(record_id, record)
    creators = " | ".join(meta.get("creators") or [])
    if "Smiljanov" not in creators:
        raise RuntimeError(f"Record {record_id} creator guard failed: {creators}")
    entries = get_file_entries(session, record)
    chosen = choose_source_files(entries)
    if not chosen:
        raise RuntimeError(f"Record {record_id} has no suitable PDF/DOCX/TXT/MD source file")

    record_objects: list[dict[str, Any]] = []
    extracted_files: list[dict[str, Any]] = []
    record_slug = slugify(meta.get("title") or record_id, 70)

    for source in chosen:
        log(f"DOWNLOAD {record_id} :: {source.name}")
        data, sha256 = download_file(session, source)
        text = extract_text(source.name, data)
        if len(text) < MIN_EXTRACTED_CHARS:
            log(f"WARN low extracted text for {source.name}: {len(text)} chars")
            continue
        chunks = chunk_text(text)
        file_slug = slugify(Path(source.name).stem, 60)
        for idx, chunk in enumerate(chunks, start=1):
            header = chunk_document_header(meta, source, sha256, idx, len(chunks))
            body = (header + chunk + "\n").encode("utf-8")
            key = (
                f"{base_prefix}/02_smiljanov_papers_chunks/zenodo/"
                f"{record_id}-{record_slug}/{file_slug}/chunk-{idx:04d}.md"
            )
            r2.put(key, body, "text/markdown; charset=utf-8")
            record_objects.append({"key": key, "bytes": len(body), "sha256": hashlib.sha256(body).hexdigest()})
        extracted_files.append({
            "name": source.name,
            "source_url": source.url,
            "zenodo_checksum": source.checksum,
            "download_sha256": sha256,
            "source_bytes": len(data),
            "extracted_chars": len(text),
            "chunk_count": len(chunks),
        })

    if not record_objects:
        raise RuntimeError(f"Record {record_id} produced zero AI-ready chunks")

    completion = dict(meta)
    completion.update({
        "synced_at": datetime.now(timezone.utc).isoformat(),
        "policy": "additive-no-delete",
        "target_namespace": "02_smiljanov_papers_chunks/zenodo",
        "source_files": extracted_files,
        "objects": record_objects,
        "object_count": len(record_objects),
    })
    marker_body = (json.dumps(completion, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode("utf-8")
    r2.put(marker, marker_body, "application/json; charset=utf-8")
    record_objects.append({"key": marker, "bytes": len(marker_body), "sha256": hashlib.sha256(marker_body).hexdigest()})
    log(f"OK {record_id}: {len(record_objects)-1} chunks + marker")
    return {
        "record_id": record_id,
        "title": meta.get("title"),
        "status": "uploaded",
        "objects": len(record_objects),
        "chunks": len(record_objects) - 1,
        "marker": marker,
    }


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--record-ids", required=True, help="Comma-separated Zenodo record IDs")
    p.add_argument("--bucket", default=os.environ.get("R2_BUCKET", DEFAULT_BUCKET))
    p.add_argument("--prefix", default=os.environ.get("TARGET_PREFIX", DEFAULT_PREFIX))
    p.add_argument("--report", default="zenodo-sync-report.json")
    args = p.parse_args()

    account_id = os.environ.get("CLOUDFLARE_ACCOUNT_ID") or os.environ.get("CF_ACCOUNT_ID")
    api_token = os.environ.get("CLOUDFLARE_API_TOKEN") or os.environ.get("CF_API_TOKEN")
    if not account_id or not api_token:
        raise SystemExit("Missing CLOUDFLARE_ACCOUNT_ID/CF_ACCOUNT_ID or CLOUDFLARE_API_TOKEN/CF_API_TOKEN")

    record_ids = [x.strip() for x in args.record_ids.split(",") if x.strip()]
    if not record_ids or any(not re.fullmatch(r"\d+", x) for x in record_ids):
        raise SystemExit("Invalid --record-ids")

    session = requests.Session()
    session.headers.update({"User-Agent": USER_AGENT})
    r2 = R2Client(account_id, api_token, args.bucket, session)
    existing = r2.list_keys(args.prefix + "/")
    log(f"R2 inventory: {len(existing)} objects under {args.prefix}/")

    results: list[dict[str, Any]] = []
    failures: list[dict[str, str]] = []
    for record_id in record_ids:
        try:
            results.append(sync_record(session, r2, args.prefix, record_id, existing))
        except Exception as exc:
            log(f"ERROR {record_id}: {exc}")
            failures.append({"record_id": record_id, "error": str(exc)})

    manifest = {
        "schema_version": "1.0",
        "operation": "WPA Zenodo remaining-after-WP004 to R2 AI-ready sync",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "bucket": args.bucket,
        "prefix": args.prefix,
        "policy": "additive-no-delete",
        "record_ids": record_ids,
        "requested_records": len(record_ids),
        "uploaded_records": sum(1 for r in results if r.get("status") == "uploaded"),
        "skipped_existing_records": sum(1 for r in results if r.get("status") == "skipped_existing"),
        "failures": failures,
        "results": results,
    }
    report_path = Path(args.report)
    report_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")

    if failures:
        log(json.dumps(manifest, ensure_ascii=False, indent=2))
        return 2

    manifest_key = f"{args.prefix}/_metadata/zenodo/remaining-after-wp004-sync.json"
    manifest_body = report_path.read_bytes()
    r2.put(manifest_key, manifest_body, "application/json; charset=utf-8")

    # Verification: all completion markers + manifest must be visible in a fresh R2 listing.
    fresh = r2.list_keys(args.prefix + "/")
    expected_markers = [f"{args.prefix}/_metadata/zenodo/{rid}.json" for rid in record_ids]
    missing = [k for k in [*expected_markers, manifest_key] if k not in fresh]
    if missing:
        raise RuntimeError(f"R2 verification failed; missing objects: {missing}")

    manifest["verified_at"] = datetime.now(timezone.utc).isoformat()
    manifest["verified_object_count_under_prefix"] = len(fresh)
    report_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    # Update manifest with verification fields.
    r2.put(manifest_key, report_path.read_bytes(), "application/json; charset=utf-8")
    log(
        f"SYNC COMPLETE: uploaded={manifest['uploaded_records']} "
        f"skipped={manifest['skipped_existing_records']} failures=0; "
        f"verified R2 objects={len(fresh)}"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
